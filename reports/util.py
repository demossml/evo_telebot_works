from bd.model import (
    Shop,
    Products,
    Documents,
    Employees,
    Session,
    Plan,
    GroupUuidAks,
    CashRegister,
    Status,
    Сonsent,
)
from arrow import utcnow, get
from typing import List, Tuple
from pprint import pprint

from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

from collections import defaultdict
from pprint import pprint
import time
from decimal import Decimal
from collections import OrderedDict

from openpyxl import Workbook
from openpyxl.utils import get_column_letter
import sys

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging

logger = logging.getLogger(__name__)

# Принимает словарь с данными о продукте


def format_sell_groups(_dict: dict, since: str, until: str) -> list[dict]:
    """
    :param _dict: словарь с данными о продукте

    :return: [
    {
        '1 Наименование:': str,
        '2 Остаток:': str,
        '3 Цена поставки:': str,
        '4 Цена продажи:': str,
        '5 Сумма(цена поставки):': str
    }

    ]
    """
    result = []
    cost_price = 0
    for k, v in _dict.items():
        prod = Products.objects(uuid=k, group__exact=False).first()
        result.append(
            {
                "1 Наименование:": prod.name,
                "2 Остаток:": "{} {}.".format(v["col"], prod.measureName),
                "3 Цена поставки: ": "{} ₱".format(prod.costPrice),
                "4 Цена продажи:": "{} ₱".format(prod.price),
                "5 Сумма(цена поставки):": "{} ₱".format(v["sum"]),
            }
        )
        cost_price += prod.costPrice
    result.append(
        {
            "⬇️⬇️⬇️⬇️Итого⬇️⬇️⬇️⬇️".upper(): " ",
            "Цена продажи:".upper(): f"{cost_price} ₱",
            "Начало периода:".upper(): since[0:10],
            "Окончание периода:".upper(): until[0:10],
        }
    )

    return result


def get_products(session: Session, shop_id: str) -> object:
    """
    :param session: Объект сессии пользователя
    :param shop_id: Список идентификаторов магазинов (UUID)
    :return: Объект с данными по продуктам магазина
    """
    # Определить роль сотрудника в сессии
    if session.employee.role == "CASHIER":
        # Если роль сотрудника - кассир, вернуть продукты с фильтрами
        # для кассиров, исключая продукты с указанным UUID
        return Products.objects(shop_id__in=shop_id, group__exact=True, uuid__in=None)
    elif session.employee.role == "ADMIN":
        # Если роль сотрудника - администратор, вернуть продукты с фильтрами
        # для администраторов магазина
        return Products.objects(shop_id__in=shop_id, group__exact=True)
    else:
        # Если роль сотрудника не определена или не соответствует "CASHIER" или "ADMIN",
        # можно добавить обработку других ролей или выбросить исключение
        raise ValueError("Недопустимая роль сотрудника в сессии")


def get_products_shops(session: Session, shop_id: list) -> object:
    """
    :param session:
    :param shop_id: uuid: str магазина
    :return: Данные по продуктам магазина
    """
    if session.employee.role == "CASHIER":
        return Products.objects(shop_id__in=shop_id, group__exact=True, uuid__in=None)
    if session.employee.role == "ADMIN":
        return Products.objects(shop_id__in=shop_id, group__exact=True)


def get_group(session: Session) -> dict[str:str]:
    """
    :param session:
    :return: {uuid: name} группы
    """
    shops = get_shops(session)
    shop_id = shops["shop_id"]

    group = {}
    for element in shop_id:
        for item in Products.objects(shop_id__exact=element, group__exact=True):
            if item["uuid"] not in group:
                group.update({item["uuid"]: item["name"]})
    return group


def status_shop(shop_id: str) -> bool:
    # Получаем объект статуса из базы данных для указанного магазина со статусом "deleted"
    doc_status = Status.objects(shop=shop_id, status="deleted").first()
    # Возвращаем True, если объект не найден или его статус "restore", иначе возвращаем False
    return not (doc_status and doc_status.status != "restore")


def status_employee(uuid: str) -> bool:
    # Получаем объект статуса из базы данных для указанного магазина со статусом "deleted"
    doc_status = Status.objects(employee=uuid, status="deleted").first()
    # Возвращаем True, если объект не найден или его статус "restore", иначе возвращаем False
    return not (doc_status and doc_status.status != "restore")


def get_shops(session: Session) -> dict:
    """
    :param session:
    :return: {
                'shop_id': ['shop_id', ...],
                'shop_name': 'name'
            }
    """
    # Получение параметров из сессии
    params = session.params["inputs"]["0"]
    pprint(params["shop"])

    # Создание списка для хранения уникальных идентификаторов магазинов (uuid)
    uuid = []
    # Поиск сотрудников с заданной фамилией (session.user_id) и итерация по их магазинам
    for item in Employees.objects(lastName=str(session.user_id)):
        for store in item.stores:
            if store not in uuid:
                uuid.append(store)
    # Проверка наличия параметра "shop" в запросе
    if "shop" in params:
        if params["shop"] == "all":
            # Возвращаем информацию о всех магазинах
            return {
                "shop_id": [
                    item["uuid"]  # Выбираем uuid из каждого элемента
                    for item in Shop.objects(uuid__in=uuid)
                    if status_shop(
                        item["uuid"]
                    )  # Проверяем статус магазина с помощью функции status_shop
                ],
                "shop_name": "Все".upper(),
            }
        else:
            # Возвращаем информацию о конкретном магазине, указанном в параметре "shop"
            shop = Shop.objects(uuid__exact=params["shop"]).only("name").first()
            return {"shop_id": [params["shop"]], "shop_name": shop.name}

    else:
        # Если параметр "shop" отсутствует, возвращаем информацию о всех магазинах
        return {
            "shop_id": [
                item["uuid"]
                for item in Shop.objects(
                    uuid__in=uuid
                )  # Перебираем магазины с указанными UUID
                if status_shop(
                    item["uuid"]
                )  # Проверяем статус магазина с помощью функции status_shop
            ],
            "shop_name": "Все".upper(),
        }


def get_shops_last_room(session: Session) -> dict:
    """
    :param session:
    :return: {
                'shop_id': ['shop_id', ...],
                'shop_name': 'name'
            }
    """
    # Получение параметров из сессии

    room = session["room"]
    params = session.params["inputs"][room]

    # Создание списка для хранения уникальных идентификаторов магазинов (uuid)
    uuid = []
    # Поиск сотрудников с заданной фамилией (session.user_id) и итерация по их магазинам
    for item in Employees.objects(lastName=str(session.user_id)):
        for store in item.stores:
            if store not in uuid:
                uuid.append(store)
    # Проверка наличия параметра "shop" в запросе
    if "shop" in params:
        if params["shop"] == "all":
            # Возвращаем информацию о всех магазинах
            return {
                "shop_id": [
                    item["uuid"]
                    for item in Shop.objects(uuid__in=uuid)
                    if status_shop(
                        item["uuid"]
                    )  # Перебираем магазины с указанными UUID
                ],
                "shop_name": "Все".upper(),
            }
        else:
            # Возвращаем информацию о конкретном магазине, указанном в параметре "shop"
            shop = Shop.objects(uuid__exact=params["shop"]).only("name").first()
            return {"shop_id": [params["shop"]], "shop_name": shop.name}

    else:
        # Если параметр "shop" отсутствует, возвращаем информацию о всех магазинах
        return {
            "shop_id": [
                item["uuid"]
                for item in Shop.objects(uuid__in=uuid)
                if status_shop(item["uuid"])
            ],
            "shop_name": "Все".upper(),
        }


def get_shops_user_id(session: Session) -> object:
    """
    Получить магазины, связанные с пользователем Telegram бота.

    :param session: Объект сессии пользователя.
    :return: Магазины, связанные с пользователем.
    """
    # Создаем пустой список для хранения уникальных идентификаторов магазинов (uuid).
    uuid = []
    # Итерируемся по сотрудникам с фамилией, равной user_id из сессии
    for item in Employees.objects(lastName=str(session.user_id)):
        # Для каждого найденного сотрудника итерируемся по его магазинам.
        for store in item.stores:
            if status_shop(store):
                # Проверяем, не находится ли магазин уже в списке UUID.
                if store not in uuid:
                    # Добавляем магазина в список UUID, если его там нет.
                    uuid.append(store)
    # Возвращаем объект Shop, в котором UUID магазина содержатся в списке UUID
    return Shop.objects(uuid__in=uuid)


def get_shops_all() -> object:
    """
    Получить магазины, связанные с пользователем Telegram бота.

    :param session: Объект сессии пользователя.
    :return: Магазины, связанные с пользователем.
    """
    # Создаем пустой список для хранения уникальных идентификаторов магазинов (uuid).
    uuid = []
    # Итерируемся по сотрудникам с фамилией, равной user_id из сессии
    for store in Shop.objects():
        # Для каждого найденного сотрудника итерируемся по его магазинам.
        if status_shop(store["uuid"]):
            # Проверяем, не находится ли магазин уже в списке UUID.
            if store["uuid"] not in uuid:
                # Добавляем магазина в список UUID, если его там нет.
                uuid.append(store["uuid"])
    # Возвращаем объект Shop, в котором UUID магазина содержатся в списке UUID
    return uuid


def get_shops_uuid_user_id(session: Session) -> list:
    """
    :param session: Сессия пользователя
    :return: Список уникальных UUID магазинов, связанных с пользователем
    """
    # Создаем пустой список для хранения уникальных UUID магазинов
    uuid = []
    # Итерируемся по всем сотрудникам с фамилией, совпадающей с идентификатором пользователя
    for item in Employees.objects(lastName=str(session.user_id)):
        # Для каждого сотрудника итерируемся по списку магазинов, к которым он привязан
        for store in item.stores:
            if status_shop(store):
                # Проверяем, не был ли магазин уже добавлен в список UUID
                if store not in uuid:
                    # Если магазина еще нет в списке, добавляем его
                    uuid.append(store)
    # Возвращаем список уникальных UUID магазинов
    return uuid


# Определение состояния сеанса:
# 1. Проверка наличия элементов в списке id_
# 2. Получение списка магазинов в зависимости от состояния сеанса и параметров
# 3. Формирование и возврат списка магазинов
def get_shops_in(session: Session, _in=[], id_=[]) -> object:
    uuid = []
    uuid_id = (
        "20220501-11CA-40E0-8031-49EADC90D1C4",
        "20220501-DDCF-409A-8022-486441F27458",
        "20220501-9ADF-402C-8012-FB88547F6222",
        "20220501-3254-40E5-809E-AC6BB204D373",
        "20220501-CB2E-4020-808C-E3FD3CB1A1D4",
        "20220501-4D25-40AD-80DA-77FAE02A007E",
        "20220430-A472-40B8-8077-2EE96318B7E7",
        "20220506-AE5B-40BA-805B-D8DDBD408F24",
    )

    if len(id_) > 0:
        # pprint(1)
        users = [uuid_id]
    else:
        # pprint(2)
        users = [
            element.stores
            for element in Employees.objects(lastName=str(session.user_id))
        ]

    # pprint(users)
    for i in users:
        for e in i:
            if e in _in:
                if e not in uuid:
                    # pprint(e)
                    uuid.append(e)
            if session.user_id == 490899906:
                for el in _in:
                    if el not in uuid:
                        uuid.append(el)
    # pprint(uuid)
    return Shop.objects(uuid__in=uuid)


def period_to_date_(
    session: Session,
) -> utcnow:
    """
    :param period: строка, представляющая период времени ("day", "week", "fortnight", "month", "two months", "6 months", "12 months", "24 months", "48 months")
    :return: строка с датой и временем в ISO формате, отстоящая на указанный период времени назад от текущего времени в UTC с временем 03:00.
    """
    period = session.params["inputs"]["0"]["period"]
    documents_open_session = Documents.objects(
        __raw__={
            # "openUserUuid": {"$in": user_uuid},
            "x_type": "OPEN_SESSION",
        }
    ).first()
    # Проверяем значение параметра "period" и выполняем соответствующее смещение времени
    if period == "day":
        return utcnow().to("local").replace(hour=3, minute=00).isoformat()
    if period == "week":
        return (
            utcnow().to("local").shift(days=-7).replace(hour=3, minute=00).isoformat()
        )
    if period == "fortnight":
        return (
            utcnow().to("local").shift(days=-14).replace(hour=3, minute=00).isoformat()
        )
    if period == "month":
        return (
            utcnow().to("local").shift(months=-1).replace(hour=3, minute=00).isoformat()
        )
    if period == "two months":
        return (
            utcnow().to("local").shift(months=-2).replace(hour=3, minute=00).isoformat()
        )
    if period == "6 months":
        return (
            utcnow().to("local").shift(months=-6).replace(hour=3, minute=00).isoformat()
        )
    if period == "12 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-12)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    if period == "24 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-24)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    if period == "48 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-48)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    raise Exception("Period is not supported")


def period_to_date(period: str) -> utcnow:
    """
    :param period: строка, представляющая период времени ("day", "week", "fortnight", "month", "two months", "6 months", "12 months", "24 months", "48 months")
    :return: строка с датой и временем в ISO формате, отстоящая на указанный период времени назад от текущего времени в UTC с временем 03:00.
    """
    # Проверяем значение параметра "period" и выполняем соответствующее смещение времени
    if period == "day":
        return utcnow().to("local").replace(hour=3, minute=00).isoformat()
    if period == "week":
        return (
            utcnow().to("local").shift(days=-7).replace(hour=3, minute=00).isoformat()
        )
    if period == "fortnight":
        return (
            utcnow().to("local").shift(days=-14).replace(hour=3, minute=00).isoformat()
        )
    if period == "month":
        return (
            utcnow().to("local").shift(months=-1).replace(hour=3, minute=00).isoformat()
        )
    if period == "two months":
        return (
            utcnow().to("local").shift(months=-2).replace(hour=3, minute=00).isoformat()
        )
    if period == "6 months":
        return (
            utcnow().to("local").shift(months=-6).replace(hour=3, minute=00).isoformat()
        )
    if period == "12 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-12)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    if period == "24 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-24)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    if period == "48 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-48)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    raise Exception("Period is not supported")


def period_first_day_of_the_month(period: str) -> utcnow:
    """
    :param period: строка, представляющая период времени ("day", "week", "fortnight", "month", "two months", "6 months", "12 months", "24 months", "48 months")
    :return: строка с датой и временем в ISO формате, отстоящая на указанный период времени назад от текущего времени в UTC с временем 03:00.
    """
    # Проверяем значение параметра "period" и выполняем соответствующее смещение времени
    if period == "day":
        return utcnow().to("local").replace(hour=3, minute=00).isoformat()
    if period == "week":
        return (
            utcnow().to("local").shift(days=-7).replace(hour=3, minute=00).isoformat()
        )
    if period == "fortnight":
        return (
            utcnow().to("local").shift(days=-14).replace(hour=3, minute=00).isoformat()
        )
    if period == "month":
        return (
            utcnow()
            .to("local")
            .shift(months=-1)
            .replace(day=1, hour=3, minute=00)
            .isoformat()
        )
    if period == "two months":
        return (
            utcnow()
            .to("local")
            .shift(months=-2)
            .replace(day=1, hour=3, minute=00)
            .isoformat()
        )
    if period == "6 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-6)
            .replace(day=1, hour=3, minute=00)
            .isoformat()
        )
    if period == "12 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-12)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    if period == "24 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-24)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    if period == "48 months":
        return (
            utcnow()
            .to("local")
            .shift(months=-48)
            .replace(hour=3, minute=00)
            .isoformat()
        )
    raise Exception("Period is not supported")


def period_to_date_2(period: str) -> utcnow:
    """
    :param period: day, week,  fortnight, month, two months,
    :return: utcnow + period
    """
    if period == "day":
        return utcnow().replace(hour=3, minute=00).isoformat()
    if period == "week":
        return utcnow().shift(days=7).replace(hour=3, minute=00).isoformat()
    if period == "fortnight":
        return utcnow().shift(days=14).replace(hour=3, minute=00).isoformat()
    if period == "month":
        return utcnow().shift(months=1).replace(hour=3, minute=00).isoformat()
    if period == "two months":
        return utcnow().shift(months=2).replace(hour=3, minute=00).isoformat()
    if period == "6 months":
        return utcnow().shift(months=6).replace(hour=3, minute=00).isoformat()
    if period == "12 months":
        return utcnow().shift(months=12).replace(hour=3, minute=00).isoformat()
    if period == "24 months":
        return utcnow().shift(months=24).replace(hour=3, minute=00).isoformat()
    if period == "48 months":
        return utcnow().shift(months=48).replace(hour=3, minute=00).isoformat()
    raise Exception("Period is not supported")


# Получить интервалы дат между минимальной и максимальной датами с определенным шагом в указанных единицах измерения (unit: days, weeks, fortnights, months).
# Возвращает список кортежей (минимальная дата, максимальная дата).
def get_intervals(
    min_date: str, max_date: str, unit: str, measure: float
) -> List[Tuple[str, str]]:
    """
    :param min_date: дата начала пириода
    :param max_date: дата окончания пириода
    :param unit: days, weeks,  fortnights, months
    :param measure: int шаг
    :return: List[Tuple[min_date, max_date]]
    """
    output = []
    while min_date < max_date:
        # Получить новую дату, добавив или вычтя указанное количество времени в указанных единицах измерения
        temp = get(min_date).shift(**{unit: measure}).isoformat()
        # Записать пару дат (min_date, min(temp, max_date)) в список выходных данных
        output.append((min_date, min(temp, max_date)))
        # Обновить min_date на новую дату temp
        min_date = temp
    return output


def get_employees(session: Session) -> object:
    """
    Получает список сотрудников пользователя Telegram бота на основе их session.user_id.

    :param session: Сеанс пользователя бота.
    :return: Объект (список или другая структура), содержащий информацию о сотрудниках.
    """
    # Инициализируем пустой список uuid для сохранения идентификаторов сотрудников
    uuid = []
    # Получаем список хранилищ (stores) сотрудников с фамилией, соответствующей session.user_id
    users = [
        element.stores for element in Employees.objects(lastName=str(session.user_id))
    ]
    # Итерируемся по списку хранилищ и добавляем их uuid в список uuid
    for i in users:
        for e in i:
            uuid.append(e)
    # Возвращаем  структуру, содержащую информацию о сотрудниках

    return Employees.objects(stores__in=uuid)


def get_employees(last_name: str) -> list:
    """
    Возвращает список UUID пользователей Telegram Bot, чья фамилия совпадает с заданной.

    :param last_name: Фамилия сотрудника для поиска.
    :return: Список UUID пользователей Telegram Bot.
    """
    # Используем список-выражение для получения UUID сотрудников, у которых фамилия совпадает с заданной.
    users = [element.uuid for element in Employees.objects(lastName=last_name)]

    # Возвращаем список UUID пользователей.
    return users


def get_products_all(session: Session, shop_id: list) -> dict[str:str]:
    """
    Получает список продуктов для указанных магазинов.

    :param shop_id: Список идентификаторов магазинов, для которых нужно получить продукты.
    :param session: Сессия (предположительно, объект сессии для взаимодействия с базой данных).
    :return: Словарь, содержащий информацию о продуктах в формате {'uuid': 'name'}.
    """
    # Создаем пустой словарь для хранения информации о продуктах
    result = {}
    for item in Products.objects(shop_id__in=shop_id, group__exact=True):
        # Если 'uuid' продукта еще не добавлен в результат, то добавляем его в словарь.
        if item["uuid"] not in result:
            # Добавляем информацию о продукте в словар
            result[item["uuid"]] = item["name"]
    # Возвращаем словарь с информацией о продуктах
    return result


def get_period(session: Session) -> dict[str:str]:
    """
    Функция get_period_day принимает объект сессии Session и возвращает словарь,
    содержащий информацию о временном периоде 'since' и 'until'.

    :param session:
    :return: {'since': str, 'until': str}
    """
    # Список возможных периодов
    period_in = ["day", "week", "fortnight", "month"]
    try:
        # Проверка, что указанный период находится в списке возможных периодов
        if session.params["inputs"]["0"]["period"] not in period_in:
            # Если период - "day", то 'since' устанавливаем в начало дня, а 'until' в конец дня
            return {
                "since": get(session.params["inputs"]["0"]["openDate"])
                .replace(day=1)
                .isoformat(),
                "until": get(session.params["inputs"]["0"]["openDate"])
                .ceil("month")
                .isoformat(),
            }
        # Если выбран период "day", возвращаем начальную дату как период "сегодня", а конечную - текущую дату и время
        if session.params["inputs"]["0"]["period"] == "day":
            return {
                "since": period_to_date(session.params["inputs"]["0"]["period"]),
                "until": utcnow().isoformat(),
            }
        # Если выбран другой период, возвращаем начальную и конечную дату с учетом указанных дат
        else:
            return {
                "since": get(session.params["inputs"]["0"]["openDate"])
                .replace(hour=3, minute=00)
                .isoformat(),
                "until": get(session.params["inputs"]["0"]["closeDate"])
                .replace(hour=23, minute=00)
                .isoformat(),
            }
    except Exception as e:
        logger.exception("Error in get_period function")
        logger.error(f"Ошибка: {e} на строке {sys.exc_info()[-1].tb_lineno}")
        raise  # Re-raise the exception after logging


def get_period_(session: Session) -> dict[str:str]:
    """
    Функция get_period_day принимает объект сессии Session и возвращает словарь,
    содержащий информацию о временном периоде 'since' и 'until'.

    :param session:
    :return: {'since': str, 'until': str}
    """
    room = session["room"]
    # pprint(room)
    # pprint(session.params["inputs"][room]["period"])
    # Список возможных периодов
    period_in = ("day", "week", "fortnight", "month")
    # Проверка, что указанный период находится в списке возможных периодов
    if session.params["inputs"][room]["period"] not in period_in:
        # Если период - "day", то 'since' устанавливаем в начало дня, а 'until' в конец дня
        return {
            "since": get(session.params["inputs"][room]["openDate"])
            .replace(day=1)
            .isoformat(),
            "until": get(session.params["inputs"][room]["openDate"])
            .ceil("month")
            .isoformat(),
        }
    # Если выбран период "day", возвращаем начальную дату как период "сегодня", а конечную - текущую дату и время
    if session.params["inputs"][room]["period"] == "day":
        return {
            "since": period_to_date(session.params["inputs"][room]["period"]),
            "until": utcnow().isoformat(),
        }
    # Если выбран другой период, возвращаем начальную и конечную дату с учетом указанных дат
    else:
        return {
            "since": get(session.params["inputs"][room]["openDate"])
            .replace(hour=3, minute=00)
            .isoformat(),
            "until": get(session.params["inputs"][room]["closeDate"])
            .replace(hour=23, minute=00)
            .isoformat(),
        }


def get_period_day(session: Session) -> dict[str:str]:
    """
    :param session:
    :return: {'since': str, 'until': str}
    """
    # Проверяем, является ли период "day"
    if session.params["inputs"]["0"]["period"] == "day":
        # Если период - "day", то возвращаем период с начала дня до конца дня
        return {
            "since": period_to_date(session.params["inputs"]["0"]["period"]),
            "until": get(period_to_date(session.params["inputs"]["0"]["period"]))
            .replace(hour=23, minute=59)
            .isoformat(),
        }

    else:
        # Если период не "day", то возвращаем весь день (от 00:01 до 23:59) указанной даты
        return {
            "since": get(session.params["inputs"]["0"]["openDate"])
            .floor("day")
            .isoformat(),
            "until": get(session.params["inputs"]["0"]["openDate"])
            .ceil("day")
            .isoformat(),
        }


def get_period_order(session: Session) -> dict[str:str]:
    """
    :param session:
    :return: {'since': str, 'until': str}
    """
    # Получаем данные о периоде заказа
    data = get(period_to_date(session.params["inputs"]["0"]["period"]))
    # pprint(data)
    if session.params["inputs"]["0"]["period"] == "day":
        # Строим диапазон "since" с началом дня и концом дня, сдвинутыми на 7 дней назад.
        return {
            "since": get(period_to_date(session.params["inputs"]["0"]["period"]))
            .shift(days=-7)
            .replace(hour=00, minute=1)
            .isoformat(),
            # Строим диапазон "since" с началом дня и концом дня, сдвинутыми на 7 дней назад.
            "until": utcnow().shift(days=-7).replace(hour=23, minute=59).isoformat(),
        }

    else:
        # В противном случае возвращаем период на основе даты открытия
        return {
            "since": get(session.params["inputs"]["0"]["openDate"])
            .shift(days=-7)
            .replace(hour=0, minute=1)
            .isoformat(),
            "until": get(session.params["inputs"]["0"]["openDate"])
            .shift(days=-7)
            .replace(hour=23, minute=59)
            .isoformat(),
        }


def get_commodity_balances(session: Session) -> dict[str:int]:
    """
    :param session:
    :return: {'uuid': str, 'quantity': init}
    """
    # Создаем пустой словарь для хранения балансов товаров
    commodity_balances = {}
    # Получаем параметры из сессии
    params = session.params["inputs"]["0"]
    # Список типов транзакций
    x_type = ("SELL", "PAYBACK", "ACCEPT")
    # Получаем список магазинов из функции get_shops
    shops = get_shops(session)
    shop_id = shops["shop_id"]
    # Итерируемся по идентификаторам магазинов
    for shop_uuid in shop_id:
        # Если параметр "group" равен "all", получаем все товары в магазине
        if params["group"] == "all":
            products = Products.objects(
                __raw__={
                    "shop_id": shop_uuid,
                }
            )

        else:
            # Иначе получаем товары, принадлежащие определенной группе
            products = Products.objects(
                __raw__={"shop_id": shop_uuid, "parentUuid": params["group"]}
            )
        # Получаем идентификаторы всех товаров в текущей группе
        products_uuid = [element.uuid for element in products]
        # Итерируемся по идентификаторам товаров
        for uuid in products_uuid:
            # Ищем документы, связанные с текущим товаром
            documents = (
                Documents.objects(
                    __raw__={
                        "shop_id": shop_uuid,
                        "x_type": {"$in": x_type},
                        "transactions.commodityUuid": uuid,
                    }
                )
                .order_by("-closeDate")
                .first()
            )

            if documents:
                for trans in documents["transactions"]:
                    if trans["x_type"] == "REGISTER_POSITION":
                        if trans["commodityUuid"] == uuid:
                            if documents.x_type == "SELL":
                                # Обновляем баланс товара при продаже
                                commodity_balances[trans["commodityUuid"]] = (
                                    trans["balanceQuantity"] - trans["quantity"]
                                )

                            if documents.x_type == "PAYBACK":
                                # Обновляем баланс товара при возврате
                                commodity_balances[trans["commodityUuid"]] = (
                                    trans["balanceQuantity"] + trans["quantity"]
                                )

                            if documents.x_type == "ACCEPT":
                                # Обновляем баланс товара при приемке
                                commodity_balances[trans["commodityUuid"]] = trans[
                                    "balanceQuantity"
                                ]
    # Возвращаем словарь с балансами товаров
    return commodity_balances


def get_commodity_balances_all(shop_id, products_uuid) -> dict[str:int]:
    """
    :param session:
    :return: {'uuid': str, 'quantity': init}
    """

    # Создаем пустой словарь для хранения балансов товаров
    x_type = ("SELL", "PAYBACK", "ACCEPT")

    quantity_data = {}

    for uuid in products_uuid:
        quantity = 0
        documents = (
            Documents.objects(
                __raw__={
                    "shop_id": shop_id,
                    "x_type": {"$in": x_type},
                    "transactions.commodityUuid": uuid,
                }
            )
            .order_by("-closeDate")
            .first()
        )
        if documents:
            for trans in documents["transactions"]:
                if trans["x_type"] == "REGISTER_POSITION":
                    if trans["commodityUuid"] == uuid:
                        if documents.x_type == "SELL":
                            # Обновляем баланс товара при продаже
                            quantity = trans["balanceQuantity"] - trans["quantity"]

                        if documents.x_type == "PAYBACK":
                            # Обновляем баланс товара при возврате
                            quantity = trans["balanceQuantity"] + trans["quantity"]

                        if documents.x_type == "ACCEPT":
                            # Обновляем баланс товара при приемке
                            quantity = trans["balanceQuantity"]
        # Возвращаем словарь с балансами товаров
        quantity_data.update({uuid: quantity})
    return quantity_data


def generate_plan():
    """
    Эта функция создает и возвращает планы для магазинов на основе определенных условий.
    """
    # Список идентификаторов групп
    group_id = (
        "78ddfd78-dc52-11e8-b970-ccb0da458b5a",
        "bc9e7e4c-fdac-11ea-aaf2-2cf05d04be1d",
        "0627db0b-4e39-11ec-ab27-2cf05d04be1d",
        "2b8eb6b4-92ea-11ee-ab93-2cf05d04be1d",
        "8a8fcb5f-9582-11ee-ab93-2cf05d04be1d",
        "97d6fa81-84b1-11ea-b9bb-70c94e4ebe6a",
        "ad8afa41-737d-11ea-b9b9-70c94e4ebe6a",
        "568905bd-9460-11ee-9ef4-be8fe126e7b9",
        "568905be-9460-11ee-9ef4-be8fe126e7b9",
    )

    shops = Shop.objects()

    _dict = {}

    for shop in shops:
        period = [7, 14, 21, 28]
        for element in period:
            # Рассчитываем временные интервалы с использованием UTC времени
            since = utcnow().shift(days=-element).replace(hour=3, minute=00).isoformat()
            # pprint(since)
            until = (
                utcnow().shift(days=-element).replace(hour=21, minute=00).isoformat()
            )
            # Получаем продукты для магазина и группы товаров
            products = Products.objects(
                __raw__={"shop_id": shop["uuid"], "parentUuid": {"$in": group_id}}
            )

            products_uuid = [element.uuid for element in products]
            # pprint(products_uuid)
            # Определяем типы транзакций
            x_type = ("SELL", "PAYBACK")
            # Получаем документы из базы данных на основе фильтров
            documents = Documents.objects(
                __raw__={
                    "closeDate": {"$gte": since, "$lt": until},
                    "shop_id": shop["uuid"],
                    "x_type": {"$in": x_type},
                    "transactions.commodityUuid": {"$in": products_uuid},
                }
            )

            sum_sell = 0
            # Вычисляем и добавляем результат в словарь
            if len(documents) > 0:
                sum_sell = 0
                for doc in documents:
                    for trans in doc["transactions"]:
                        if trans["x_type"] == "REGISTER_POSITION":
                            if trans["commodityUuid"] in products_uuid:
                                sum_sell += trans["sum"]
            else:
                sum_sell = 0
            # pprint(sum_sell)

            _day = 4
            # Обновляем словарь суммами продаж
            if shop["uuid"] in _dict:
                _dict[shop["uuid"]] += round(
                    (int(sum_sell) / _day + int(sum_sell) / 5 / 100 * 5)
                )
            else:
                _dict[shop["uuid"]] = round(
                    (int(sum_sell) / _day + int(sum_sell) / 5 / 100 * 5)
                )
    # Обновляем параметры плана в базе данных
    for k, v in _dict.items():
        params = {"closeDate": utcnow().isoformat(), "shop_id": k}
        # Выставляем минимальное значение 3500
        params["sum"] = max(int(5200), v)

        # print(params)
        Plan.objects(closeDate=utcnow().isoformat()).update(**params, upsert=True)
    # Возвращаем параметры плана
    return params


def remainder(shops_uuid: list) -> list[dict[str:str]]:
    x_type = ("SELL", "PAYBACK", "ACCEPT")

    result = {}
    for sh in shops_uuid:
        result[sh] = []
        shop = Shop.objects(uuid=sh).only("name").first()

        products = Products.objects(__raw__={"shop_id": sh, "group": False})
        # pprint(products)
        uuids = [i.uuid for i in products]
        # pprint(uuids)
        for uuid in uuids:
            product = Products.objects(uuid=uuid).first()
            # pprint(product)
            cost_price = product["costPrice"]
            # pprint(cost_price)
            documents = (
                Documents.objects(
                    __raw__={
                        "shop_id": sh,
                        "x_type": {"$in": x_type},
                        "transactions.commodityUuid": uuid,
                    }
                )
                .order_by("-closeDate")
                .first()
            )

            # pprint(documents)
            # pprint('e')
            if documents is not None:
                for trans in documents["transactions"]:
                    if trans["x_type"] == "REGISTER_POSITION":
                        if trans["commodityUuid"] == uuid:
                            if documents.x_type == "SELL":
                                if trans["balanceQuantity"] - trans["quantity"] > 0:
                                    result[sh].append(
                                        {
                                            trans["commodityUuid"]: trans[
                                                "balanceQuantity"
                                            ]
                                            - trans["quantity"]
                                        }
                                    )
                            if documents.x_type == "PAYBACK":
                                # pprint(trans)
                                if trans["balanceQuantity"] + trans["quantity"] > 0:
                                    result[sh].append(
                                        {
                                            trans["commodityUuid"]: trans[
                                                "balanceQuantity"
                                            ]
                                            + trans["quantity"]
                                        }
                                    )
                            if documents.x_type == "ACCEPT":
                                # pprint(documents.closeDate)
                                result[sh].append(
                                    {trans["commodityUuid"]: trans["quantity"]}
                                )
    return result


# Зп по назначенным группам аксессуаров
def get_aks_salary(shop_id: str, since_: str, until_: str) -> dict:
    """
    Возвращает информацию о зарплате сотрудников в отделе aks.

    Args:
        shop_id (str): Идентификатор магазина.
        since_ (str): Начальная дата для анализа продаж.
        until_ (str): Конечная дата для анализа продаж.

    Returns:
        dict: Словарь с информацией о продажах и бонусах сотрудников aks.
            - "accessory_sum_sell" (int): Сумма продаж аксессуаров.
            - "bonus_accessory" (int): Рассчитанный бонус сотрудникам за продажи аксессуаров.
    """
    # Получаем документы aks для указанного магазина и периода времени
    documents_aks = (
        GroupUuidAks.objects(
            __raw__={
                "closeDate": {"$lte": until_[:10]},
                "shop_id": shop_id,
                "x_type": "MOTIVATION_PARENT_UUID",
            }
        )
        .order_by("-closeDate")  # Сортировка по убыванию даты закрытия
        .first()
    )
    if documents_aks:
        # Получаем группу продуктов, связанных с документами aks
        group = Products.objects(
            __raw__={
                "shop_id": shop_id,
                "parentUuid": {"$in": documents_aks.parentUuids},
            }
        )
        # Получаем список UUID продуктов
        products_uuid = [i.uuid for i in group]
        # Получаем документы о продажах для указанного магазина, периода и продуктов aks
        documents_sale = Documents.objects(
            __raw__={
                "closeDate": {"$gte": since_, "$lt": until_},
                "shop_id": shop_id,
                "x_type": "SELL",
                "transactions.commodityUuid": {"$in": products_uuid},
            }
        )
        # Вычисляем общую сумму продаж аксессуаров.
        sum_sales_aks = 0
        for doc in documents_sale:
            for trans in doc["transactions"]:
                if trans["x_type"] == "REGISTER_POSITION":
                    if trans["commodityUuid"] in products_uuid:
                        sum_sales_aks += trans["sum"]
        # Вычисляем бонус за продажу аксессуаров (5% от общей суммы продаж).
        return {
            "accessory_sum_sell": sum_sales_aks,
            "bonus_accessory": round(int(sum_sales_aks / 100 * 5) / 10) * 10,
        }
    else:
        # Если не найдено документов GroupUuidAks, возвращаем нули.
        return {
            "accessory_sum_sell": 0,
            "bonus_accessory": 0,
        }


# Зп по продаже мотивационого товара
def get_mot_salary(shop_id: str, since_: str, until_: str) -> dict:
    """
    Функция возвращает бонус мотивации для магазина на основе продаж и данных о мотивации.

    Args:
        shop_id (str): Идентификатор магазина.
        since_ (str): Начальная дата для анализа продаж (включительно).
        until_ (str): Конечная дата для анализа продаж (исключительно).

    Returns:
        dict: Словарь, содержащий бонус мотивации в виде "bonus_motivation".
    """

    # Запрашиваем данные о мотивации для магазина
    documents_mot = (
        GroupUuidAks.objects(
            __raw__={
                "closeDate": {"$lte": until_[:10]},
                "shop_id": shop_id,
                "x_type": "MOTIVATION_UUID",
            }
        )
        .order_by("-closeDate")
        .first()
    )
    if documents_mot:
        # Извлекаем идентификаторы продуктов из данных о мотивации
        products_uuid = [k for k, v in documents_mot.uuid.items()]
        # Запрашиваем данные о продажах продуктов в указанный период времени
        documents_sale = Documents.objects(
            __raw__={
                "closeDate": {"$gte": since_, "$lt": until_},
                "shop_id": shop_id,
                "x_type": "SELL",
                "transactions.commodityUuid": {"$in": products_uuid},
            }
        )
        # Создаем словарь для хранения данных о зарплате
        dict_salary = {}

        for doc in documents_sale:
            for trans in doc["transactions"]:
                if trans["x_type"] == "REGISTER_POSITION":
                    if trans["commodityUuid"] in products_uuid:
                        # Если продукт уже присутствует в словаре, увеличиваем его количество
                        if trans["commodityUuid"] in dict_salary:
                            dict_salary[trans["commodityUuid"]] += trans["quantity"]
                        # Если продукта нет в словаре, добавляем его с количеством из текущей транзакции
                        else:
                            dict_salary[trans["commodityUuid"]] = trans["quantity"]
        # Инициализируем сумму мотивации
        sum_mot = 0
        # Вычисляем сумму мотивации на основе данных о продажах и мотивации
        for k, v in dict_salary.items():
            sum_mot += v * documents_mot.uuid[k]
        # Возвращаем бонус мотивации в виде словаря
        return {"bonus_motivation": int(sum_mot)}

    else:
        # Если данных о мотивации нет, возвращаем нулевой бонус мотивации
        return {"bonus_motivation": 0}


# План по группе, продажи, бонус за выполнение плана
def get_plan_bonus(shop_id: str, since_: str, until_: str) -> dict:
    """
    Функция возвращает информацию о бонусах за выполнение плана для магазина на основе данных о планах и продажах.

    Args:
        shop_id (str): Идентификатор магазина.
        since_ (str): Начальная дата периода.
        until_ (str): Конечная дата периода.

    Returns:
        dict: Словарь с информацией о бонусе:
            - "plan_motivation_prod": Заданный план продаж.
            - "sales_motivation_prod": Фактические продажи в периоде.
            - "bonus_motivation_prod": Рассчитанный бонус за выполнение плана.
    """

    # Поиск планов для магазина и периода
    plan_ = Plan.objects(
        __raw__={
            "closeDate": {"$gte": since_, "$lt": until_},
            "shop_id": shop_id,
        }
    )
    # Если нет найденных планов, генерируем новый
    if len(plan_) > 0:
        plan = Plan.objects(
            __raw__={
                "closeDate": {"$gte": since_, "$lt": until_},
                "shop_id": shop_id,
            }
        ).first()
    else:
        generate_plan()
        plan = Plan.objects(
            __raw__={
                "closeDate": {"$gte": since_, "$lt": until_},
                "shop_id": shop_id,
            }
        ).first()
    # Группы товаров для анализа
    group_id = (
        "bc9e7e4c-fdac-11ea-aaf2-2cf05d04be1d",
        "568905bd-9460-11ee-9ef4-be8fe126e7b9",
        "2b8eb6b4-92ea-11ee-ab93-2cf05d04be1d",
        "568905be-9460-11ee-9ef4-be8fe126e7b9",
        "ad8afa41-737d-11ea-b9b9-70c94e4ebe6a",
        "8a8fcb5f-9582-11ee-ab93-2cf05d04be1d",
        "78ddfd78-dc52-11e8-b970-ccb0da458b5a",
    )
    # Поиск товаров в указанных группах
    products = Products.objects(
        __raw__={"shop_id": shop_id, "parentUuid": {"$in": group_id}}
    )
    # Список UUID товаров
    products_uuid = [element.uuid for element in products]
    # Типы транзакций, которые учитываем (продажи и возвраты)
    x_type = ("SELL", "PAYBACK")
    # Получаем документы с продажами и возвратами товаров из списка products_uuid
    documents_2 = Documents.objects(
        __raw__={
            "closeDate": {"$gte": since_, "$lt": until_},
            "shop_id": shop_id,
            "x_type": {"$in": x_type},
            "transactions.commodityUuid": {"$in": products_uuid},
        }
    )

    sum_sell_today = 0
    for doc_2 in documents_2:
        for trans_2 in doc_2["transactions"]:
            if trans_2["x_type"] == "REGISTER_POSITION":
                if trans_2["commodityUuid"] in products_uuid:
                    sum_sell_today += trans_2["sum"]
    # Получаем информацию о мотивации из последнего документа
    documents_motiv = (
        GroupUuidAks.objects(
            __raw__={
                "closeDate": {"$lte": until_[:10]},
                "shop_id": shop_id,
                "x_type": "MOTIVATION",
            }
        )
        .order_by("-closeDate")
        .first()
    )
    # Функция для расчета бонуса по выполнению плана
    if documents_motiv:
        if sum_sell_today >= plan.sum:
            return {
                "plan_motivation_prod": plan.sum,
                "sales_motivation_prod": sum_sell_today,
                "bonus_motivation_prod": documents_motiv.motivation,
            }
        else:
            return {
                "plan_motivation_prod": plan.sum,
                "sales_motivation_prod": sum_sell_today,
                "bonus_motivation_prod": 0,
            }

    else:
        return {
            "plan_motivation_prod": plan.sum,
            "sales_motivation_prod": sum_sell_today,
            "bonus_motivation_prod": 0,
        }


# Функция для получения оклада продавцов в магазине
def get_salary(shop_id: str, until_: str) -> dict:
    """
    Возвращает оклад продавцов в магазине до указанной даты.

    Args:
        shop_id (str): Идентификатор магазина.
        until_ (str): Дата, до которой нужно получить оклад (включительно).

    Returns:
        dict: Словарь с информацией о окладе. Если информация найдена, то содержит "salary",
              иначе "salary" равно 0.
    """

    # Ищем документы о зарплате продавцов в указанном магазине до указанной даты
    documents_salary = GroupUuidAks.objects(
        __raw__={
            "closeDate": {"$lte": until_[:10]},
            "shop_id": shop_id,
            "x_type": "SALARY",
        }
    ).first()
    # Если найдена информация о зарплате, возвращаем ее
    if documents_salary:
        return {
            "salary": documents_salary.salary,
        }
    else:
        # Если документ не найден, возвращаем нулевую зарплату (0)
        return {
            "salary": 0,
        }


# Функция для получения доплаты к окладу продавцов на основе их идентификатора и даты окончания
def get_surcharge(employee_uuid: str, until_: str) -> dict:
    """
    Функция для получения доплаты к окладу продавцов.

    Args:
        employee_uuid (str): Уникальный идентификатор сотрудника.
        until_ (str): Дата (включительно), до которой нужно получить доплату.

    Returns:
        dict: Словарь, содержащий информацию о доплате к окладу:
            - "surcharge" (int): Размер доплаты, если есть, или 0, если доплаты нет.
    """
    # pprint(employee_uuid)
    # Ищем документы о доплатах к окладу для указанного сотрудника и даты окончания
    documents_surcharge = (
        GroupUuidAks.objects(
            __raw__={
                "closeDate": {"$lte": until_[:10]},
                "employee_uuid": employee_uuid,
                "x_type": "ASSING_A_SURCHARGE",
            }
        )
        .order_by("-closeDate")
        .first()
    )
    # pprint(documents_surcharge)
    # Если есть информация о доплате, возвращаем её
    if documents_surcharge:
        return {
            "surcharge": documents_surcharge.surcharge,
        }
    else:
        # Если информации нет, возвращаем 0
        return {
            "surcharge": 0,
        }


def get_total_salary(
    employee_uuid: str, shop_id: str, since_: str, until_: str
) -> dict:
    """
    Возвращает структуру, содержащую информацию о заработной плате сотрудника.

    Args:
        employee_uuid (str): Уникальный идентификатор сотрудника.
        shop_id (str): Идентификатор магазина.
        since_ (str): Дата начала периода.
        until_ (str): Дата окончания периода.

    Returns:
        dict: Словарь с информацией о заработной плате сотрудника.

    Структура словаря:
    {
        "accessory_sum_sell": int,         # Сумма продаж аксессуаров
        "bonus_accessory": int,           # Бонус за продажу аксессуаров
        "bonus_motivation": int,          # Бонус за мотивацию
        "plan_motivation_prod": int,      # План мотивации продаж
        "sales_motivation_prod": int,     # Продажи мотивации продаж
        "bonus_motivation_prod": int,     # Бонус за мотивацию продаж
        "salary": int,                   # Заработная плата
        "surcharge": int,                # Надбавка
        "total_salary": int,             # Общая заработная плата
        "closeDate": until_[:10]         # Дата окончания периода (обрезанная до дня)
    }
    """
    # Инициализируем пустой словарь, в который будем добавлять информацию о заработке сотрудника
    result = {}
    # Получаем информацию о заработной плате с помощью различных функций
    result.update(get_aks_salary(shop_id, since_, until_))
    pprint(1)
    result.update(get_mot_salary(shop_id, since_, until_))
    pprint(2)
    result.update(get_plan_bonus(shop_id, since_, until_))
    pprint(3)
    result.update(get_salary(shop_id, until_))
    pprint(4)
    result.update(get_surcharge(employee_uuid, until_))
    pprint(5)
    # Добавляем дату окончания периода в результат (обрезанную до дня)
    result.update({"closeDate": until_})
    pprint(7)
    # Вычисляем общую заработную плату
    total_salary = (
        result["bonus_accessory"]
        + result["bonus_motivation"]
        + result["bonus_motivation_prod"]
        + result["salary"]
        + result["surcharge"]
    )
    # Добавляем общую заработную плату в результат
    result.update({"total_salary": total_salary})
    # Возвращаем словарь result с информацией о заработке сотрудника
    return result


# Функция для сбора статистики по продажам {uuid: commodity}
def gather_statistics_uuid(documents, products_uuid):
    """
    Собирает статистику по продажам для каждого товара в заданных документах.

    Parameters:
    - documents: Список документов продажи
    - products_uuid: Список UUID товаров

    Returns:
    - data_sale: Словарь с количеством проданных товаров для каждого UUID товара
    """
    data_sale = {}
    for doc in documents:
        for trans in doc["transactions"]:
            if trans["x_type"] == "REGISTER_POSITION":
                if trans["commodityUuid"] in products_uuid:
                    if trans["commodityUuid"] in data_sale:
                        data_sale[trans["commodityUuid"]] += trans["quantity"]
                    else:
                        data_sale[trans["commodityUuid"]] = trans["quantity"]
    return data_sale


# Функция для сбора статистики по продажам {name: commodity}
def gather_statistics_name(documents, products_uuid):
    """
    Собирает статистику по продажам для каждого товара в заданных документах.

    Parameters:
    - documents: Список документов продажи
    - products_uuid: Список UUID товаров

    Returns:
    - data_sale: Словарь с количеством проданных товаров для каждого UUID товара
    """
    data_sale = {}
    for doc in documents:
        for trans in doc["transactions"]:
            if trans["x_type"] == "REGISTER_POSITION":
                if trans["commodityUuid"] in products_uuid:
                    if trans["commodityName"] in data_sale:
                        data_sale[trans["commodityName"]] += trans["quantity"]
                    else:
                        data_sale[trans["commodityName"]] = trans["quantity"]
    return data_sale


def cash() -> int:
    sum_ = 0
    cash = CashRegister.objects()

    if cash:
        for doc in cash:
            # pprint(doc["cash"])
            if doc["x_type"] == "CASH_INCOME":
                sum_ += int(doc["cash"])

            if doc["x_type"] == "CASH_OUTCOME":
                sum_ -= int(doc["cash"])

        return sum_
    else:
        return sum_


def last_time(shop_id: str) -> dict[str:str]:
    # Определение временного периода для анализа
    since = utcnow().replace(hour=3, minute=00).isoformat()
    until = utcnow().isoformat()

    shop = Shop.objects(uuid__exact=shop_id).only("name").first()
    Documents_last_time = (
        Documents.objects(
            __raw__={
                "closeDate": {"$gte": since, "$lt": until},
                "shop_id": shop_id,
            }
        )
        .order_by("-closeDate")
        .only("closeDate")
        .first()
    )
    if Documents_last_time:
        time = get(Documents_last_time.closeDate).shift(hours=3).isoformat()[11:19]
    else:
        time = 0

    return {f"🕰️ выг. {shop.name}": time}


def get_sale_uuid(shop_id: list[str], since: str, until: str) -> list:
    """
    Args:
        shop_id (list[str]): uuid
        since (str): iso
        until (str): iso

    Returns:
        list: uuid
    """

    data_uud = []

    intervals = get_intervals(since, until, "days", 1)

    for since_, until_ in intervals:
        documents = Documents.objects(
            __raw__={
                "closeDate": {"$gte": since_, "$lt": until_},
                "shop_id": {"$in": shop_id},
                "x_type": "SELL",
            }
        )
        for doc in documents:
            for trans in doc["transactions"]:
                if trans["x_type"] == "REGISTER_POSITION":
                    if trans["commodityUuid"] not in data_uud:
                        data_uud.append(trans["commodityUuid"])
    return data_uud


def process_uuid(
    uuid,
    shop_uuid: list,
):
    # Инициализируем баланс товара
    commodity_balance = 0

    x_type = ("SELL", "PAYBACK", "ACCEPT")

    # Получаем документы для заданного магазина, товара и типов транзакций
    documents = (
        Documents.objects(
            __raw__={
                "shop_id": shop_uuid,
                "x_type": {"$in": x_type},
                "transactions.commodityUuid": uuid,
            }
        ).order_by("-closeDate")
    ).first()

    if documents:
        # Если найдены документы, обрабатываем транзакции
        for trans in documents.transactions:
            # Проверяем тип транзакции и обновляем баланс товара
            if (
                trans["x_type"] == "REGISTER_POSITION"
                and trans["commodityUuid"] == uuid
            ):
                if documents.x_type == "SELL":
                    commodity_balance += trans["balanceQuantity"] - trans["quantity"]
                elif documents.x_type == "PAYBACK":
                    commodity_balance += trans["balanceQuantity"] + trans["quantity"]
                elif documents.x_type == "ACCEPT":
                    commodity_balance += trans["balanceQuantity"]
    # Возвращаем кортеж с uuid товара и балансом

    return uuid, commodity_balance


def get_commodity_balances_p(shop_id: list, product_uuid: list) -> defaultdict:
    #
    commodity_balances = defaultdict(int)

    # Используем ThreadPoolExecutor для создания потоков
    with ThreadPoolExecutor() as executor:
        for shop_uuid in shop_id:
            # Используем executor.map для параллельного выполнения process_uuid для каждого товара
            results = executor.map(
                lambda element: process_uuid(element, shop_uuid),
                [element for element in product_uuid],
            )

            # Обновляем словарь с балансами товаров
            for uuid, balance in results:
                commodity_balances[uuid] += balance

    return commodity_balances


def process_shop(shop_uuid, group_id, period):
    _dict = {}

    for element in period:
        since = utcnow().shift(days=-element).replace(hour=3, minute=00).isoformat()
        until = utcnow().shift(days=-element).replace(hour=21, minute=00).isoformat()

        products = Products.objects(
            __raw__={"shop_id": shop_uuid, "parentUuid": {"$in": group_id}}
        )

        products_uuid = [element.uuid for element in products]

        x_type = ("SELL", "PAYBACK")

        documents = Documents.objects(
            __raw__={
                "closeDate": {"$gte": since, "$lt": until},
                "shop_id": shop_uuid,
                "x_type": {"$in": x_type},
                "transactions.commodityUuid": {"$in": products_uuid},
            }
        )

        sum_sell = 0

        if len(documents) > 0:
            for doc in documents:
                for trans in doc["transactions"]:
                    if trans["x_type"] == "REGISTER_POSITION":
                        if trans["commodityUuid"] in products_uuid:
                            sum_sell += trans["sum"]

        # pprint(sum_sell)

        _day = 4

        if shop_uuid in _dict:
            _dict[shop_uuid] += round(
                (int(sum_sell) / _day + int(sum_sell) / 5 / 100 * 5)
            )
        else:
            _dict[shop_uuid] = round(
                (int(sum_sell) / _day + int(sum_sell) / 5 / 100 * 5)
            )

    return _dict


def generate_plan_():
    group_id = (
        "78ddfd78-dc52-11e8-b970-ccb0da458b5a",
        "bc9e7e4c-fdac-11ea-aaf2-2cf05d04be1d",
        "0627db0b-4e39-11ec-ab27-2cf05d04be1d",
        "2b8eb6b4-92ea-11ee-ab93-2cf05d04be1d",
        "8a8fcb5f-9582-11ee-ab93-2cf05d04be1d",
        "97d6fa81-84b1-11ea-b9bb-70c94e4ebe6a",
        "ad8afa41-737d-11ea-b9b9-70c94e4ebe6a",
        "568905bd-9460-11ee-9ef4-be8fe126e7b9",
        "568905be-9460-11ee-9ef4-be8fe126e7b9",
    )

    shop_id = [element.uuid for element in Shop.objects()]
    period = [7, 14, 21, 28]

    _dict = {}

    with ThreadPoolExecutor() as executor:
        # Запускаем потоки для каждого магазина
        results = list(
            executor.map(lambda shop: process_shop(shop, group_id, period), shop_id)
        )

    # Объединяем результаты из всех потоков
    for result in results:
        for shop_uuid, value in result.items():
            if shop_uuid in _dict:
                _dict[shop_uuid] += value
            else:
                _dict[shop_uuid] = value

    for k, v in _dict.items():
        params = {"closeDate": utcnow().isoformat(), "shop_id": k}

        params["sum"] = max(int(5200), v)

        Plan.objects(closeDate=utcnow().isoformat()).update(**params, upsert=True)

    return params


def get_plan(shop: str) -> object:
    since = utcnow().replace(hour=3, minute=00).isoformat()
    until = utcnow().replace(hour=20, minute=59).isoformat()

    # Получение данных о планах продаж для магазина
    plan = (
        Plan.objects(
            __raw__={
                "closeDate": {"$gte": since, "$lt": until},
                "shop_id": shop,
            }
        )
        .only("sum")
        .first()
    )
    if plan:
        return plan
    else:
        # pprint("generate_plan_")
        generate_plan_()
        return (
            Plan.objects(
                __raw__={
                    "closeDate": {"$gte": since, "$lt": until},
                    "shop_id": shop,
                }
            )
            .only("sum")
            .first()
        )


def analyze_sales_for_shop(shop_id) -> dict:
    # Группы товаров для анализа продаж
    group_id = (
        "78ddfd78-dc52-11e8-b970-ccb0da458b5a",
        "bc9e7e4c-fdac-11ea-aaf2-2cf05d04be1d",
        "0627db0b-4e39-11ec-ab27-2cf05d04be1d",
        "2b8eb6b4-92ea-11ee-ab93-2cf05d04be1d",
        "8a8fcb5f-9582-11ee-ab93-2cf05d04be1d",
        "97d6fa81-84b1-11ea-b9bb-70c94e4ebe6a",
        "ad8afa41-737d-11ea-b9b9-70c94e4ebe6a",
        "568905bd-9460-11ee-9ef4-be8fe126e7b9",
        "568905be-9460-11ee-9ef4-be8fe126e7b9",
    )
    # Определение временного периода для анализа
    since = utcnow().replace(hour=3, minute=00).isoformat()
    until = utcnow().isoformat()

    # Словарь для хранения данных о продажах по магазинам

    # dict_last_time.update(last_time(shop["uuid"]))
    since = utcnow().replace(hour=3, minute=00).isoformat()
    until = utcnow().replace(hour=20, minute=59).isoformat()

    # Получение списка продуктов, относящихся к группам товаров
    products = Products.objects(
        __raw__={"shop_id": shop_id, "parentUuid": {"$in": group_id}}
    )

    # Формирование списка идентификаторов продуктов
    products_uuid = [element.uuid for element in products]

    # Типы операций для анализа (продажи и возвраты)
    x_type = ("SELL", "PAYBACK")

    # Получение документов о продажах и возвратах для продуктов
    documents_sale = Documents.objects(
        __raw__={
            "closeDate": {"$gte": since, "$lt": until},
            "shop_id": shop_id,
            "x_type": {"$in": x_type},
            "transactions.commodityUuid": {"$in": products_uuid},
        }
    )

    sum_sell_today = 0
    # Вычисление суммы продаж за текущий период
    for doc in documents_sale:
        for trans in doc["transactions"]:
            if trans["x_type"] == "REGISTER_POSITION":
                if trans["commodityUuid"] in products_uuid:
                    sum_sell_today += trans["sum"]
    return {shop_id: sum_sell_today}

    # pprint(sales_data)


def analyze_sales_parallel():
    sales_data = defaultdict(int)

    with ThreadPoolExecutor() as executor:
        # Получаем список магазинов
        shops = get_shops_all()

        # Запускаем выполнение задачи для каждого магазина в отдельном потоке
        future_to_shop = {
            executor.submit(analyze_sales_for_shop, shop): shop for shop in shops
        }

        # Дожидаемся завершения всех потоков и собираем результаты
        for future in as_completed(future_to_shop):
            shop = future_to_shop[future]
            try:
                result = future.result()
                sales_data.update(result)
            except Exception as e:
                logger.error(f"An error occurred for shop {shop}: {e}")
                logger.error(f"Ошибка: {e} на строке {sys.exc_info()[-1].tb_lineno}")

    # pprint(sales_data)
    return sales_data


def cash_outcome_for_shop(shop_id, since, until):
    payment_category = {
        1: "Инкассация".upper(),
        2: "Оплата поставщику".upper(),
        3: "Оплата услуг".upper(),
        4: "Аренда".upper(),
        5: "Заработная плата".upper(),
        6: "Прочее".upper(),
    }

    sum_payment_category = {}

    x_type = "CASH_OUTCOME"

    documents = Documents.objects(
        __raw__={
            "closeDate": {"$gte": since, "$lt": until},
            "shop_id": shop_id,
            "x_type": x_type,
        }
    )

    for doc in documents:
        if doc["x_type"] == "CASH_OUTCOME":
            for trans in doc["transactions"]:
                if trans["x_type"] == "CASH_OUTCOME":
                    category = payment_category.get(trans["paymentCategoryId"])
                    if category:
                        sum_payment_category[category] = (
                            sum_payment_category.get(category, 0) + trans["sum"]
                        )

    return {shop_id: sum_payment_category}


def cash_outcome_parallel(shops: list, since: str, until: str) -> dict:
    sum_payment_category = defaultdict(dict)

    with ThreadPoolExecutor() as executor:
        # Запускаем потоки для каждого магазина
        results = list(
            executor.map(lambda shop: cash_outcome_for_shop(shop, since, until), shops)
        )

    # Объединяем результаты из всех потоков
    for result in results:
        for shop_id, categories in result.items():
            sum_payment_category[shop_id].update(categories)

    # pprint(dict(sum_payment_category))
    return dict(sum_payment_category)


def cash_outcome(shop_id, since, until):
    payment_category = {
        1: "Инкассация",
        2: "Оплата поставщику",
        3: "Оплата услуг",
        4: "Аренда",
        5: "Заработная плата",
        6: "Прочее",
    }

    sum_payment_category = {}

    x_type = "CASH_OUTCOME"

    documents = Documents.objects(
        __raw__={
            "closeDate": {"$gte": since, "$lt": until},
            "shop_id": shop_id,
            "x_type": x_type,
        }
    )

    for doc in documents:
        if doc["x_type"] == "CASH_OUTCOME":
            for trans in doc["transactions"]:
                if trans["x_type"] == "CASH_OUTCOME":
                    category = payment_category.get(trans["paymentCategoryId"])
                    if category:
                        sum_payment_category[category] = sum_payment_category.get(
                            category, 0
                        ) + Decimal(trans["sum"]).quantize(Decimal("0.00"))
    return sum_payment_category


def json_to_xls_format_change(
    data_list: list,
):
    # Создаем новую книгу Excel
    book = Workbook()

    # Выбираем активный лист в книге
    sheet = book.active

    # Задаем порядок столбцов
    columns_name = ["name", "sum", "average_sales", "sales_days"]

    # Записываем названия столбцов в первую строку
    for col_idx, column_name in enumerate(columns_name, start=1):
        sheet.cell(row=1, column=col_idx, value=column_name)

    # Записываем данные из входного списка в ячейки
    for row_idx, item in enumerate(data_list, start=2):
        # Проходим по названиям столбцов
        for col_idx, column_name in enumerate(columns_name, start=1):
            # Записываем значение в соответствующую ячейку, если оно есть в словаре
            sheet.cell(row=row_idx, column=col_idx, value=item.get(column_name))

    # Возвращаем созданную книгу Excel и количество удаленных дубликатов
    return book


# Функция для вычисления продаж для одного магазина
def calculate_sales(
    shop,
    group_id,
    since_,
):
    """
    Рассчитывает продажи для одного магазина в указанный временной период.

    :param shop: Информация о магазине.
    :param group_id: Список идентификаторов групп товаров.
    :param date: Дата.
    :return: Рассчитанная средняя сумма продаж для магазина.
    """

    intervals_plan = get_intervals_plan(since_)

    shop_name = Shop.objects(uuid=shop).only("name").first().name
    sum_sell = 0
    for since, until in intervals_plan:

        products = Products.objects(
            __raw__={"shop_id": shop, "parentUuid": {"$in": group_id}}
        )
        products_uuid = [element.uuid for element in products]

        x_type = ("SELL", "PAYBACK")
        # Получаем документы из базы данных на основе фильтров
        documents = Documents.objects(
            __raw__={
                "closeDate": {"$gte": since, "$lt": until},
                "shop_id": shop,
                "x_type": {"$in": x_type},
                "transactions.commodityUuid": {"$in": products_uuid},
            }
        )

        if len(documents) > 0:
            for doc in documents:
                for trans in doc["transactions"]:
                    if trans["x_type"] == "REGISTER_POSITION":
                        if trans["commodityUuid"] in products_uuid:
                            sum_sell += trans["sum"]
        else:
            sum_sell += 0
    # Рассчитываем результат с учетом минимальной суммы
    min_result = 4
    result = max(min_result, int(Decimal((sum_sell / 4) * 1.05).quantize(Decimal("0"))))

    pprint(f"{shop_name} | {sum_sell}/{result} | {since}/{until} ")
    params = {
        "closeDate": since_,
        "shop_id": shop,
        "sum": result,
    }
    Plan.objects(closeDate=utcnow().isoformat()).update(**params, upsert=True)

    return {shop_name: f"{result} | {since_[:10]}"}


# Функция для получения временных интервалов для заданной даты
def get_intervals_plan(date) -> list:
    output = []
    period = [7, 14, 21, 28]
    for element in period:
        since = get(date).shift(days=-element).replace(hour=0, minute=1).isoformat()
        until = get(date).shift(days=-element).replace(hour=23, minute=59).isoformat()
        output.append((since, until))
    return output


# Основная функция для генерации планов продаж для магазинов параллельно
def generate_plan_parallel(shops, start_date, end_date):
    group_id = (
        "78ddfd78-dc52-11e8-b970-ccb0da458b5a",
        "bc9e7e4c-fdac-11ea-aaf2-2cf05d04be1d",
        "0627db0b-4e39-11ec-ab27-2cf05d04be1d",
        "2b8eb6b4-92ea-11ee-ab93-2cf05d04be1d",
        "8a8fcb5f-9582-11ee-ab93-2cf05d04be1d",
        "97d6fa81-84b1-11ea-b9bb-70c94e4ebe6a",
        "ad8afa41-737d-11ea-b9b9-70c94e4ebe6a",
        "568905bd-9460-11ee-9ef4-be8fe126e7b9",
        "568905be-9460-11ee-9ef4-be8fe126e7b9",
    )

    # Получение временных интервалов для каждого периода
    intervals = get_intervals(start_date, end_date, "days", 1)

    result_data = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        # Список для хранения будущих результатов
        future_to_shop = {
            executor.submit(calculate_sales, shop, group_id, since): shop
            for since, until in intervals
            for shop in shops
        }
        for future in as_completed(future_to_shop):
            shop = future_to_shop[future]
            try:
                result = future.result()
                result_data.append(result)
            except Exception as e:
                print(f"Exception occurred for shop {shop}: {e}")

    # Сортировка по shop_name и since_[:10]}
    sorted_result_data = sorted(
        result_data,
        key=lambda x: (list(x.keys())[0], x[list(x.keys())[0]].split(" | ")[1]),
    )

    # pprint(sorted_result_data)
    return sorted_result_data


def get_sales_by_category(
    shop_id: str,
    since: str,
    until: str,
) -> dict:
    """
    Возвращает словарь, содержащий сумму продаж по категориям платежей за определенный период.

    Args:
        shop_id (str): Идентификатор магазина.
        since (str): Начальная дата периода (включительно).
        until (str): Конечная дата периода (исключительно).

    Returns:
        dict: Словарь с суммами продаж по категориям платежей.
              Структура: {shop_id: {"payment_type": total_sales, ...}}
    """
    # Используем defaultdict для автоматического создания ключей с начальным значением 0
    payment_type_sum_sell = defaultdict(Decimal)

    # Ищем документы, соответствующие условиям
    documents = Documents.objects(
        __raw__={
            "closeDate": {"$gte": since, "$lt": until},
            "shop_id": shop_id,
            "x_type": "SELL",
        }
    )
    # Итерируемся по найденным документам  documents (продажам)
    for doc in documents:
        # Итерируемся по транзакциям в каждом документе
        for trans in doc["transactions"]:
            # Если тип транзакции "PAYMENT" (оплата)
            if trans["x_type"] == "PAYMENT":

                # Увеличиваем сумму продаж по определенному типу платежа
                payment_type_sum_sell[trans["paymentType"]] += Decimal(
                    trans["sum"]
                ).quantize(Decimal("0.00"))

    return {shop_id: dict(payment_type_sum_sell)}


def sales_parallel(shops: list, since: str, until: str) -> dict:
    """
    Возвращает словарь, содержащий сумму продаж по категориям платежей для нескольких магазинов.

    Args:
        shops (list): Список идентификаторов магазинов.
        since (str): Начальная дата периода (включительно).
        until (str): Конечная дата периода (исключительно).

    Returns:
        dict: Словарь с суммами продаж по категориям платежей для каждого магазина.
              Структура: {shop_id: {"payment_type": total_sales, ...}, ...}
    """
    sales_results = {}

    with ThreadPoolExecutor() as executor:
        # Запускаем обработку каждого магазина в отдельном потоке
        futures = {
            executor.submit(get_sales_by_category, shop_id, since, until): shop_id
            for shop_id in shops
        }

        # Получаем результаты вычислений
        for future in futures:
            shop_id = futures[future]
            try:
                result = future.result()
                sales_results.update(result)
            except Exception as e:
                print(f"Error processing shop {shop_id}: {e}")

    return sales_results


# Функция для обработки данных по кассе для одного магазина
def get_cash(shop_id):

    # Инициализация структуры для хранения данных по кассе для данного магазина
    report_data = defaultdict(int)

    # Получение последнего документа типа FPRINT_Z_REPORT для данного магазина
    document_fprint = (
        Documents.objects(
            __raw__={
                "shop_id": shop_id,
                "x_type": "FPRINT",
                "transactions.x_type": "FPRINT_Z_REPORT",
            },
        )
        .order_by("-closeDate")
        .first()
    )

    # Начальные параметры для фильтрации документов
    documents_query = {"shop_id": shop_id}

    # Если найден документ FPRINT_Z_REPORT, обновляем данные по кассе
    if document_fprint:

        # Обновление данных по кассе с учетом последнего FPRINT_Z_REPORT
        for trans in document_fprint.transactions:
            report_data[shop_id] += Decimal(trans["cash"]).quantize(Decimal("0.00"))
        since = document_fprint.closeDate
        until = utcnow().isoformat()
        documents_query["closeDate"] = {"$gte": since, "$lt": until}
    # Получение всех документов для данного магазина
    documents = Documents.objects(__raw__=documents_query)

    # Обработка каждого документа и его транзакций
    for doc in documents:
        amount = Decimal(trans.get("sum", 0)).quantize(Decimal("0.00"))

        if doc["x_type"] == "CASH_OUTCOME":
            for trans in doc["transactions"]:
                if trans["x_type"] == "CASH_OUTCOME":
                    report_data[shop_id] -= amount

        if doc["x_type"] == "CASH_INCOME":
            for trans in doc["transactions"]:
                if trans["x_type"] == "CASH_INCOME":
                    report_data[shop_id] += amount
        if doc["x_type"] == "SELL":
            for trans in doc["transactions"]:
                if trans["x_type"] == "PAYMENT":
                    if trans["paymentType"] == "CASH":
                        report_data[shop_id] += amount
        if doc["x_type"] == "PAYBACK":
            for trans in doc["transactions"]:
                if trans["x_type"] == "PAYMENT":
                    if trans["paymentType"] == "CASH":
                        report_data[shop_id] -= amount
    return dict(report_data)


def calculate_for_shops(shops):
    result_data = {}
    with ThreadPoolExecutor() as executor:
        # Параллельное выполнение функции для каждого магазина
        future_to_shop = {executor.submit(get_cash, shop): shop for shop in shops}
        for future in as_completed(future_to_shop):
            shop = future_to_shop[future]
            try:
                result_data.update(future.result())
            except Exception as e:
                print(f"Error processing shop {shop}: {e}")
    return result_data


s = "2024-02-21T10:12:23.000+0000"
s = "2024-02-21T21:00:00.000+0000"


def get_top_n_sales(sales_by_product: dict, n=50):
    """
    Функция для получения топ N элементов из словаря с продажами по продуктам.

    :param sales_by_product: Исходный словарь с продажами по продуктам.
    :param n: Количество элементов для включения в топ (по умолчанию 50).
    :return: Словарь с топ N элементами по продажам.
    """
    # Сортировка словаря по значениям в порядке убывания
    sorted_sales = dict(
        OrderedDict(sorted(sales_by_product.items(), key=lambda t: -t[1]))
    )

    # Выбор только первых N элементов
    top_n_sales = dict(list(sorted_sales.items())[:n])

    # Возврат результата
    return top_n_sales, len(top_n_sales)


def xls_to_json_format_change(book):

    # Получаем активный лист из книги Excel
    ws = book.active

    my_list = []  # Создаем пустой список для хранения словарей

    # Находим номер последнего столбца и строки
    last_column = len(list(ws.columns))
    last_row = len(list(ws.rows))

    # Проходимся по каждой строке в таблице Excel
    for row in range(1, last_row + 1):
        my_dict = {}  # Создаем пустой словарь для текущей строки
        # Проходимся по каждому столбцу в текущей строке
        for column in range(1, last_column + 1):
            column_letter = get_column_letter(
                column
            )  # Получаем буквенное обозначение столбца
            if row > 1:  # Пропускаем первую строку, так как это заголовки
                # Добавляем элементы в словарь в формате "значение заголовка: значение ячейки"
                my_dict[ws[column_letter + str(1)].value] = ws[
                    column_letter + str(row)
                ].value
        if len(my_dict) > 0:  # Убеждаемся, что словарь не пустой
            my_list.append(my_dict)  # Добавляем словарь в список
    return my_list  # Возвращаем список словарей


import arrow
from arrow import get


def calculate_difference(open_date: str, status_open_date: str):
    try:
        # Извлекаем время из строки open_date (например, "2024-10-10 12:30:00")
        open_time_str = open_date[11:16]  # Предполагаем, что время в формате HH:mm
        time_open_date = get(open_time_str, "HH:mm")

        if status_open_date:
            # Конвертируем status_open_date также в объект arrow
            time_status_open_date = get(status_open_date, "HH:mm")

            # Вычисляем разницу во времени в минутах
            time_difference = (
                time_open_date - time_status_open_date
            ).total_seconds() / 60

            if time_difference > 0:
                return time_difference
            else:
                return None
        else:
            # Если status_open_date не предоставлен (или является None/False)
            return 0
    except Exception as e:
        # Обработка ошибок для диагностики
        print(f"Ошибка при расчёте: {e}")
        return None


# Анкета/Questionnaire


# Функция для создания директории, если она не существует
def create_directory_if_not_exists(directory):
    if not os.path.exists(directory):
        os.makedirs(directory)


# Функция для создания документа на основе входных данных
def create_document(data):
    try:
        # Функция для добавления заголовка блока
        def add_block_title(doc, title):
            title_paragraph = doc.add_heading(title, level=2)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Функция для добавления подзаголовка
        def add_sub_title(doc, title):
            sub_title_paragraph = doc.add_paragraph(title)
            sub_title_paragraph.style.font.bold = True

        # Создаем новый документ
        doc = Document()
        doc.add_heading("АНКЕТА", level=1)

        # Заполнение документа данными из JSON

        # Блок личной информации
        add_block_title(doc, "Личная информация")
        doc.add_paragraph(f'Фамилия: {data["full_name"].split()[0]}')
        doc.add_paragraph(f'Имя: {data["full_name"].split()[1]}')
        doc.add_paragraph(f'Отчество: {data["full_name"].split()[2]}')
        doc.add_paragraph(f'Дата рождения: {data["dateOf_birth"]}')
        doc.add_paragraph(f'Гражданство: {data["citizenship"]}')
        doc.add_paragraph(
            f'Место рождения (село город край область республика): {data["place_of_birth"]}'
        )
        doc.add_paragraph(f'Адрес (место жительства): {data["residence_address"]}')
        doc.add_paragraph(f'Адрес (место прописки): {data["registration_address"]}')
        doc.add_paragraph(f'Домашний телефон: {data["home_phone"]}')
        doc.add_paragraph("Сотовый телефон: [Not provided]")
        doc.add_paragraph("Рабочий телефон: [Not provided]")

        # Блок паспортных данных
        add_block_title(doc, "Паспортные данные")
        doc.add_paragraph("Паспортные данные: [Not provided]")

        # Блок семейного положения
        add_block_title(doc, "Семейное положение")
        doc.add_paragraph(f'Семейное положение: {data["counterparty"]}')
        doc.add_paragraph("Дети (пол возраст): [Not provided]")

        # Блок сведений о близких родственниках
        add_block_title(doc, "Сведения о близких родственниках")
        for relative in data["relatives_information"]:
            add_sub_title(doc, f'Степень родства: {relative["close_relati"]}')
            doc.add_paragraph(f'Ф.И.О.: {relative["full_name"]}')
            doc.add_paragraph(f'Дата рождения: {relative["dateOf_birth"]}')
            doc.add_paragraph(
                f'Место работы: {relative.get("сompany", "[Not provided]")}'
            )
            doc.add_paragraph("Должность: [Not provided]")
            doc.add_paragraph("Телефон: [Not provided]")
            doc.add_paragraph("Адрес (место жительства): [Not provided]")

        # Блок образования
        add_block_title(doc, "Образование")
        for edu in data["education"]:
            add_sub_title(doc, "Образование:")
            doc.add_paragraph(f'Дата поступления: {edu["education_start_date"]}')
            doc.add_paragraph(f'Дата окончания: {edu["education_end_date"]}')
            doc.add_paragraph(
                f'Название учебного заведения: {edu["education_institution_name"]}'
            )
            doc.add_paragraph(f'Специальность: {edu["specialization"]}')

        # Блок дополнительного образования
        add_block_title(doc, "Дополнительное образование")
        # Если нет отдельного ключа для дополнительного образования, используем тот же, что и для основного образования
        for edu in data["education"]:
            add_sub_title(doc, "Дополнительное образование:")
            doc.add_paragraph(f'Дата поступления: {edu["education_start_date"]}')
            doc.add_paragraph(f'Дата окончания: {edu["education_end_date"]}')
            doc.add_paragraph(
                f'Название учебного заведения: {edu["education_institution_name"]}'
            )
            doc.add_paragraph(f'Специальность: {edu["specialization"]}')

        # Блок навыков и умений
        add_block_title(doc, "Навыки и умения")
        doc.add_paragraph(
            f'Навыки владения компьютером с какими программными продуктами приходилось работать: {data.get("skills", "[Not provided]")}'
        )
        doc.add_paragraph("Знание иностранных языков степень владения: [Not provided]")

        # Блок рекомендателей
        add_block_title(doc, "Рекомендатели")
        doc.add_paragraph(
            "Рекомендатели (должность Ф.И.О. и контактный телефон): [Not provided]"
        )

        # Блок трудовой деятельности
        add_block_title(doc, "Трудовая деятельность")
        doc.add_paragraph(
            "Трудовая деятельность (укажите в обратном хронологическом порядке 5 последних мест Вашей работы):"
        )
        for work in data["works_information"]:
            doc.add_paragraph(f'Дата начала: {work["work_start_date"]}')
            doc.add_paragraph(f'Дата окончания: {work["work_end_date"]}')
            doc.add_paragraph(f'Наименование организации: {work["company"]}')
            doc.add_paragraph(f'Должность: {work["position"]}')
            doc.add_paragraph("Адрес организации: [Not provided]")
            doc.add_paragraph(
                f'Причина увольнения (фактическая): {work["reason_forLeaving"]}'
            )

        # Блок дополнительной информации
        add_block_title(doc, "Дополнительная информация")
        doc.add_paragraph(f'Желаемый уровень заработной платы: {data["desiredSalary"]}')
        doc.add_paragraph(f'Преимущества Вашей кандидатуры: {data["advantages"]}')
        doc.add_paragraph(f'Ваши хобби: {data["hobbies"]}')
        doc.add_paragraph(
            "Какую информацию Вы хотели бы добавить о себе: [Not provided]"
        )

        # Блок согласия на проверку информации
        add_block_title(doc, "Согласие на проверку информации")
        doc.add_paragraph(
            "Против проверки предоставленной мною информации не возражаю."
        )

        # Блок даты и подписи
        add_block_title(doc, "Дата и подпись")
        doc.add_paragraph("Дата заполнения: [Not provided]")
        doc.add_paragraph("Подпись: [Not provided]")

        # Сохраняем документ в поток байтов
        byte_stream = BytesIO()
        byte_stream.seek(0)

        return byte_stream
    except Exception as e:
        logger.error(f"Ошибка: {e} на строке {sys.exc_info()[-1].tb_lineno}")


def generate_text_message(user_id):
    try:
        data = Сonsent.objects(user_id=int(user_id)).first()

        if not data:
            raise ValueError("No data found for the given user_id.")

        not_provided = "[Not provided] ❌"
        message = ""

        # Функция для добавления заголовка блока
        def add_block_title(title, emoji):
            nonlocal message
            message += f"\n\n{emoji} **{title}**\n"

        # Функция для добавления подзаголовка
        def add_sub_title(title):
            nonlocal message
            message += f"\n{title}\n"

        # Заголовок анкеты
        message += "📝 **АНКЕТА**\n"

        # Заполнение документа данными из JSON

        # Блок личной информации
        add_block_title("Личная информация", "👤")
        full_name = data.full_name.split()
        message += (
            f"Фамилия: {full_name[0] if len(full_name) > 0 else not_provided}\n"
            f"Имя: {full_name[1] if len(full_name) > 1 else not_provided}\n"
            f"Отчество: {full_name[2] if len(full_name) > 2 else not_provided}\n"
            f'Дата рождения: {getattr(data, "dateOf_birth", not_provided)}\n'
            f'Гражданство: {getattr(data, "citizenship", not_provided)}\n'
            f'Место рождения (село город край область республика): {getattr(data, "place_of_birth", not_provided)}\n'
            f'Адрес (место жительства): {getattr(data, "residence_address", not_provided)}\n'
            f'Адрес (место прописки): {getattr(data, "registration_address", not_provided)}\n'
            f'Домашний телефон: {getattr(data, "home_phone", not_provided)}\n'
            f"Сотовый телефон: {not_provided}\n"
            f"Рабочий телефон: {not_provided}\n"
        )

        # Блок паспортных данных
        add_block_title("Паспортные данные", "🛂")
        message += f"Паспортные данные: {not_provided}\n"

        # Блок семейного положения
        add_block_title("Семейное положение", "👪")
        message += f'Семейное положение: {getattr(data, "counterparty", not_provided)}\nДети (пол возраст): {not_provided}\n'

        # Блок сведений о близких родственниках
        add_block_title("Сведения о близких родственниках", "👨‍👩‍👧‍👦")
        for relative in getattr(data, "relatives_information", []):
            add_sub_title(
                f'Степень родства: {relative.get("close_relati", not_provided)}'
            )
            message += (
                f'Ф.И.О.: {relative.get("full_name", not_provided)}\n'
                f'Дата рождения: {relative.get("dateOf_birth", not_provided)}\n'
                f'Место работы: {relative.get("сompany", not_provided)}\n'
                f"Должность: {not_provided}\n"
                f"Телефон: {not_provided}\n"
                f"Адрес (место жительства): {not_provided}\n"
            )

        # Блок образования
        add_block_title("Образование", "🎓")
        for edu in getattr(data, "education", []):
            add_sub_title("Образование:")
            message += (
                f'Дата поступления: {edu.get("education_start_date", not_provided)}\n'
                f'Дата окончания: {edu.get("education_end_date", not_provided)}\n'
                f'Название учебного заведения: {edu.get("education_institution_name", not_provided)}\n'
                f'Специальность: {edu.get("specialization", not_provided)}\n'
            )

        # Блок дополнительного образования
        add_block_title("Дополнительное образование", "📚")
        for edu in getattr(data, "education", []):
            add_sub_title("Дополнительное образование:")
            message += (
                f'Дата поступления: {edu.get("education_start_date", not_provided)}\n'
                f'Дата окончания: {edu.get("education_end_date", not_provided)}\n'
                f'Название учебного заведения: {edu.get("education_institution_name", not_provided)}\n'
                f'Специальность: {edu.get("specialization", not_provided)}\n'
            )

        # Блок навыков и умений
        add_block_title("Навыки и умения", "💼")
        message += f'Навыки: {getattr(data, "skills", not_provided)}\nЗнание иностранных языков степень владения: {not_provided}\n'

        # Блок трудовой деятельности
        add_block_title("Трудовая деятельность", "🏢")
        message += "Трудовая деятельность (укажите в обратном хронологическом порядке 5 последних мест Вашей работы):\n"
        for work in getattr(data, "works_information", []):
            message += (
                f'Дата начала: {work.get("work_start_date", not_provided)}\n'
                f'Дата окончания: {work.get("work_end_date", not_provided)}\n'
                f'Наименование организации: {work.get("company", not_provided)}\n'
                f'Должность: {work.get("position", not_provided)}\n'
                f"Адрес организации: {not_provided}\n"
                f'Причина увольнения (фактическая): {work.get("reason_forLeaving", not_provided)}\n'
            )

        # Блок дополнительной информации
        add_block_title("Дополнительная информация", "ℹ️")
        message += (
            f'Желаемый уровень заработной платы: {getattr(data, "desiredSalary", not_provided)}\n'
            f'Преимущества Вашей кандидатуры: {getattr(data, "advantages", not_provided)}\n'
            f'Ваши хобби: {getattr(data, "hobbies", not_provided)}\n'
            f"Какую информацию Вы хотели бы добавить о себе: {not_provided}\n"
        )

        return message
    except Exception as e:
        logger.error(f"Ошибка: {e} на строке {sys.exc_info()[-1].tb_lineno}")
        return None
